from flask import Flask, render_template, jsonify, request, send_from_directory, send_file
import json
import tweepy
import os
import secrets
from datetime import datetime, timedelta, timezone
import sqlite3
import logging
import threading
import time

app = Flask(__name__)
app.secret_key = os.environ.get('SECRET_KEY', secrets.token_hex(32))

# 配置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Twitter API配置
TWITTER_BEARER_TOKEN = os.environ.get('TWITTER_BEARER_TOKEN')
print(f"🔑 TWITTER_BEARER_TOKEN: {TWITTER_BEARER_TOKEN[:20] if TWITTER_BEARER_TOKEN else 'None'}...")

# 数据库文件路径
DB_FILE = 'research_platform.db'

def format_interval(seconds):
    """将秒数格式化为人性化的时间显示"""
    if seconds < 3600:
        return f"{seconds // 60}分钟"
    elif seconds < 86400:
        hours = seconds // 3600
        minutes = (seconds % 3600) // 60
        if minutes == 0:
            return f"{hours}小时"
        else:
            return f"{hours}小时{minutes}分钟"
    else:
        days = seconds // 86400
        hours = (seconds % 86400) // 3600
        if hours == 0:
            return f"{days}天"
        else:
            return f"{days}天{hours}小时"

def insert_researcher_batch(cursor, batch_data, error_details):
    """批量插入研究者数据"""
    added_count = 0

    for data in batch_data:
        try:
            cursor.execute('''
                INSERT OR REPLACE INTO researchers
                (rank, name, country, company, research_focus, x_account)
                VALUES (?, ?, ?, ?, ?, ?)
            ''', data)
            added_count += 1

        except Exception as e:
            error_msg = f"插入数据失败 {data[1]}: {str(e)}"
            error_details.append(error_msg)
            logger.error(error_msg)

    return added_count

# 完整的TwitterAPI类 - 在你的app.py中找到class TwitterAPI并替换为以下完整版本：

class TwitterAPI:
    def __init__(self):
        self.client = None
        self.api_working = False
        self.connection_tested = False

        if TWITTER_BEARER_TOKEN:
            try:
                self.client = tweepy.Client(
                    bearer_token=TWITTER_BEARER_TOKEN,
                    wait_on_rate_limit=False  # 避免启动时阻塞
                )
                logger.info("✅ Twitter API客户端初始化成功")
                logger.info("⏳ Twitter API连接将在首次使用时测试")
            except Exception as e:
                logger.error(f"❌ Twitter API初始化失败: {e}")
                self.client = None
        else:
            self.client = None
            logger.warning("⚠️ Twitter Bearer Token未配置，将无法获取真实数据")

def test_connection(self):
    """测试API连接 - 修复版本"""
    if self.connection_tested:
        return self.api_working

    try:
        if self.client:
            # 使用支持Bearer Token的API调用测试
            user = self.client.get_user(username='twitter')  # 测试Twitter官方账号
            if user and user.data:
                logger.info("🔗 Twitter API连接测试成功")
                self.api_working = True
                self.connection_tested = True
                return True
            else:
                logger.error("❌ Twitter API测试失败")
                self.api_working = False
                self.connection_tested = True
        else:
            self.api_working = False
            self.connection_tested = True
    except Exception as e:
        logger.error(f"❌ API连接测试失败: {e}")
        self.api_working = False
        self.connection_tested = True
    return self.api_working

    def ensure_connection(self):
        """确保API连接可用，如果未测试则先测试"""
        if not self.connection_tested:
            self.test_connection()
        return self.api_working

    def get_user_info(self, username):
        """获取用户基本信息，包括关注者数量"""
        if not self.client:
            logger.warning(f"Twitter客户端未配置，无法获取 {username} 的用户信息")
            return None

        # 确保连接可用
        if not self.ensure_connection():
            logger.warning(f"Twitter API连接不可用，跳过获取 {username} 的用户信息")
            return None

        try:
            username = username.replace('@', '').strip()
            if not username:
                logger.warning("用户名为空")
                return None

            logger.info(f"🔍 正在获取用户信息: {username}")

            user_response = self.client.get_user(
                username=username,
                user_fields=['public_metrics', 'profile_image_url', 'description', 'verified']
            )

            if not user_response or not user_response.data:
                logger.warning(f"❌ Twitter用户 {username} 不存在或无法访问")
                return None

            user = user_response.data
            public_metrics = getattr(user, 'public_metrics', {})

            user_info = {
                'id': str(user.id),
                'username': user.username,
                'name': user.name,
                'followers_count': public_metrics.get('followers_count', 0),
                'following_count': public_metrics.get('following_count', 0),
                'tweet_count': public_metrics.get('tweet_count', 0),
                'listed_count': public_metrics.get('listed_count', 0),
                'profile_image_url': getattr(user, 'profile_image_url', ''),
                'description': getattr(user, 'description', ''),
                'verified': getattr(user, 'verified', False)
            }

            logger.info(f"✅ 成功获取 {username} 的信息: {user_info['followers_count']} 关注者, {user_info['following_count']} 正在关注")
            return user_info

        except tweepy.Unauthorized:
            logger.error(f"❌ 无权访问用户 {username}，可能是私人账户或API权限不足")
            return None
        except tweepy.NotFound:
            logger.error(f"❌ 用户 {username} 不存在")
            return None
        except tweepy.TooManyRequests:
            logger.error(f"❌ API请求过于频繁，请稍后再试")
            self.api_working = False
            return None
        except Exception as e:
            logger.error(f"❌ 获取 {username} 用户信息时发生错误: {e}")
            return None

class TwitterAPI:
    # ... 现有的 __init__, test_connection, ensure_connection, get_user_info 方法 ...
    
    # 在类的最后添加这个方法：
    def get_user_tweets(self, username, max_results=10, start_time=None, end_time=None):
        """获取用户推文"""
        if not self.client:
            return []
            
        try:
            username = username.replace('@', '').strip()
            user_response = self.client.get_user(username=username)
            if not user_response or not user_response.data:
                return []
                
            user_id = user_response.data.id
            tweets_response = self.client.get_users_tweets(
                id=user_id,
                max_results=min(max_results, 50),
                tweet_fields=['created_at', 'public_metrics']
            )
            
            if not tweets_response or not tweets_response.data:
                return []
                
            result = []
            for tweet in tweets_response.data:
                public_metrics = getattr(tweet, 'public_metrics', {})
                result.append({
                    'id': str(tweet.id),
                    'content': tweet.text or '',
                    'created_at': tweet.created_at.isoformat() if tweet.created_at else None,
                    'likes': public_metrics.get('like_count', 0),
                    'retweets': public_metrics.get('retweet_count', 0),
                    'replies': public_metrics.get('reply_count', 0),
                    'quotes': public_metrics.get('quote_count', 0),
                    'author': username,
                    'type': 'original',
                    'media_urls': [],
                    'is_retweet': False,
                    'is_reply': False
                })
            return result
            
        except Exception as e:
            logger.error(f"获取推文失败: {e}")
            return []

        except tweepy.Unauthorized:
            logger.error(f"❌ 无权访问用户 {username} 的推文，可能是私人账户")
            return []
        except tweepy.NotFound:
            logger.error(f"❌ 推文获取404错误: {username}")
            return []
        except tweepy.TooManyRequests:
            logger.error(f"❌ API请求过于频繁，请稍后再试")
            self.api_working = False
            return []
        except Exception as e:
            logger.error(f"❌ 获取 {username} 推文时发生意外错误: {e}")
            logger.error(f"   错误类型: {type(e).__name__}")
            return []

# 同时添加一个验证方法存在的调试端点到app.py中：
@app.route('/api/check_twitter_api_methods')
def check_twitter_api_methods():
    """检查TwitterAPI类的方法是否存在"""
    try:
        methods_info = {
            'twitter_api_exists': twitter_api is not None,
            'client_exists': twitter_api and twitter_api.client is not None,
            'available_methods': []
        }

        if twitter_api:
            # 检查方法是否存在
            method_names = ['get_user_info', 'get_user_tweets', 'test_connection', 'ensure_connection']
            for method_name in method_names:
                exists = hasattr(twitter_api, method_name)
                is_callable = exists and callable(getattr(twitter_api, method_name, None))
                methods_info['available_methods'].append({
                    'method': method_name,
                    'exists': exists,
                    'callable': is_callable
                })

        return jsonify(methods_info)

    except Exception as e:
        logger.error(f"检查TwitterAPI方法失败: {e}")
        return jsonify({'error': str(e)}), 500

# 添加一个测试方法是否工作的端点：
@app.route('/api/test_get_user_tweets/<username>')
def test_get_user_tweets_method(username):
    """直接测试get_user_tweets方法"""
    try:
        if not twitter_api:
            return jsonify({'error': 'twitter_api对象不存在'}), 500

        if not hasattr(twitter_api, 'get_user_tweets'):
            return jsonify({'error': 'get_user_tweets方法不存在'}), 500

        if not callable(getattr(twitter_api, 'get_user_tweets')):
            return jsonify({'error': 'get_user_tweets不是可调用的方法'}), 500

        # 调用方法
        tweets = twitter_api.get_user_tweets(username, max_results=3)

        return jsonify({
            'username': username,
            'method_exists': True,
            'tweets_count': len(tweets),
            'tweets_sample': tweets[:2] if tweets else [],
            'success': True
        })

    except Exception as e:
        logger.error(f"测试get_user_tweets方法失败: {e}")
        return jsonify({
            'error': str(e),
            'error_type': type(e).__name__,
            'username': username
        }), 500

class ResearcherManager:
    def __init__(self):
        self.init_database()
        self.load_sample_data()

    def init_database(self):
        """初始化数据库 - 支持大规模数据存储"""
        conn = sqlite3.connect('research_platform.db')
        cursor = conn.cursor()

        # 开启外键约束和基本优化设置
        cursor.execute("PRAGMA foreign_keys = ON;")
        cursor.execute("PRAGMA synchronous = NORMAL;")  # 平衡性能和安全性

        # 研究者表 - 优化字段类型和索引
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS researchers (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                rank INTEGER,
                name TEXT NOT NULL,
                country TEXT,
                company TEXT,
                research_focus TEXT,
                x_account TEXT,
                followers_count TEXT DEFAULT '0',
                following_count TEXT DEFAULT '0',
                avatar_url TEXT DEFAULT '',
                is_monitoring BOOLEAN DEFAULT 0,
                is_special_focus BOOLEAN DEFAULT 0,
                created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                updated_at DATETIME DEFAULT CURRENT_TIMESTAMP
            )
        ''')

        # 为高频查询字段创建索引
        try:
            cursor.execute('CREATE INDEX IF NOT EXISTS idx_researchers_rank ON researchers(rank);')
            cursor.execute('CREATE INDEX IF NOT EXISTS idx_researchers_name ON researchers(name);')
            cursor.execute('CREATE INDEX IF NOT EXISTS idx_researchers_monitoring ON researchers(is_monitoring);')
            cursor.execute('CREATE INDEX IF NOT EXISTS idx_researchers_special ON researchers(is_special_focus);')
            cursor.execute('CREATE INDEX IF NOT EXISTS idx_researchers_account ON researchers(x_account);')
        except Exception as e:
            logger.warning(f"创建索引时遇到警告: {e}")

        # 内容表 - 优化存储和索引
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS x_content (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                researcher_id INTEGER,
                tweet_id TEXT UNIQUE,
                content TEXT,
                content_type TEXT DEFAULT 'text',
                likes_count INTEGER DEFAULT 0,
                retweets_count INTEGER DEFAULT 0,
                replies_count INTEGER DEFAULT 0,
                created_at DATETIME,
                collected_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                is_historical BOOLEAN DEFAULT 0,
                media_urls TEXT,
                FOREIGN KEY (researcher_id) REFERENCES researchers (id) ON DELETE CASCADE
            )
        ''')

        # 内容表索引
        try:
            cursor.execute('CREATE INDEX IF NOT EXISTS idx_content_researcher ON x_content(researcher_id);')
            cursor.execute('CREATE INDEX IF NOT EXISTS idx_content_created ON x_content(created_at);')
            cursor.execute('CREATE INDEX IF NOT EXISTS idx_content_tweet_id ON x_content(tweet_id);')
        except Exception as e:
            logger.warning(f"创建内容表索引时遇到警告: {e}")

        # 监控任务表
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS monitoring_tasks (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                researcher_id INTEGER,
                status TEXT DEFAULT 'active',
                last_check DATETIME DEFAULT CURRENT_TIMESTAMP,
                check_interval INTEGER DEFAULT 3600,
                FOREIGN KEY (researcher_id) REFERENCES researchers (id) ON DELETE CASCADE
            )
        ''')

        # 监控任务表索引
        try:
            cursor.execute('CREATE INDEX IF NOT EXISTS idx_monitoring_researcher ON monitoring_tasks(researcher_id);')
            cursor.execute('CREATE INDEX IF NOT EXISTS idx_monitoring_status ON monitoring_tasks(status);')
        except Exception as e:
            logger.warning(f"创建监控任务表索引时遇到警告: {e}")

        # 系统设置表
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS system_settings (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                setting_key TEXT UNIQUE NOT NULL,
                setting_value TEXT NOT NULL,
                description TEXT,
                updated_at DATETIME DEFAULT CURRENT_TIMESTAMP
            )
        ''')

        # 插入默认监控周期设置（30分钟 = 1800秒）
        cursor.execute('''
            INSERT OR IGNORE INTO system_settings (setting_key, setting_value, description)
            VALUES ('monitoring_interval', '1800', '监控检查间隔（秒）')
        ''')

        # 创建元数据表用于跟踪数据库状态
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS db_metadata (
                key TEXT PRIMARY KEY,
                value TEXT,
                updated_at DATETIME DEFAULT CURRENT_TIMESTAMP
            )
        ''')

        conn.commit()
        conn.close()
        logger.info("✅ 数据库初始化完成 - 已优化支持大规模数据")

    def load_sample_data(self):
        """加载研究者示例数据 (此为应用基础数据，非动态内容)"""
        conn = sqlite3.connect('research_platform.db')
        cursor = conn.cursor()

        # 检查是否已经加载过示例数据
        cursor.execute('SELECT value FROM db_metadata WHERE key = ?', ('sample_data_loaded',))
        if cursor.fetchone():
            conn.close()
            return

        researchers_data = [
            {
                'rank': 1, 'name': 'Ilya Sutskever', 'country': 'Canada', 'company': 'SSI',
                'research_focus': 'AlexNet、Seq2seq、深度学习', 'x_account': '@ilyasut',
                'followers_count': '127K', 'following_count': '89'
            },
            {
                'rank': 2, 'name': 'Noam Shazeer', 'country': 'USA', 'company': 'Google Deepmind',
                'research_focus': '注意力机制、混合专家模型、角色AI', 'x_account': '@noamshazeer',
                'followers_count': '45K', 'following_count': '156'
            },
            {
                'rank': 3, 'name': 'Geoffrey Hinton', 'country': 'UK', 'company': 'University of Toronto',
                'research_focus': '反向传播、玻尔兹曼机、深度学习', 'x_account': '@geoffreyhinton',
                'followers_count': '234K', 'following_count': '67'
            },
            {
                'rank': 4, 'name': 'Alec Radford', 'country': 'USA', 'company': 'Thinking Machines',
                'research_focus': '生成对抗网络、GPT、CLIP', 'x_account': '@alec_radford',
                'followers_count': '89K', 'following_count': '123'
            },
            {
                'rank': 5, 'name': 'Andrej Karpathy', 'country': 'Slovakia', 'company': 'Tesla',
                'research_focus': '计算机视觉、神经网络、自动驾驶', 'x_account': '@karpathy',
                'followers_count': '512K', 'following_count': '234'
            }
        ]

        for researcher in researchers_data:
            cursor.execute('''
                INSERT OR IGNORE INTO researchers
                (rank, name, country, company, research_focus, x_account, followers_count, following_count)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            ''', (
                researcher['rank'], researcher['name'], researcher['country'],
                researcher['company'], researcher['research_focus'], researcher['x_account'],
                researcher['followers_count'], researcher['following_count']
            ))

        # 标记示例数据已加载
        cursor.execute('INSERT OR REPLACE INTO db_metadata (key, value) VALUES (?, ?)',
                      ('sample_data_loaded', 'true'))

        conn.commit()
        conn.close()
        logger.info("✅ 示例数据加载完成")

    def load_sample_data_if_empty(self):
        """仅在数据库为空时加载示例数据"""
        conn = sqlite3.connect('research_platform.db')
        cursor = conn.cursor()

        cursor.execute('SELECT COUNT(*) FROM researchers')
        count = cursor.fetchone()[0]

        if count == 0:
            # 重置加载标记
            cursor.execute('DELETE FROM db_metadata WHERE key = ?', ('sample_data_loaded',))
            conn.commit()
            conn.close()

            # 重新加载示例数据
            self.load_sample_data()
        else:
            conn.close()

# 监控任务 - 优化支持大规模监控
class MonitoringService:
    def __init__(self):
        self.running = False
        self.thread = None
        self.max_concurrent_checks = 10  # 最大并发检查数
        self.current_interval = self.get_monitoring_interval()  # 从数据库获取间隔

    def get_monitoring_interval(self):
        """从数据库获取监控间隔设置"""
        try:
            conn = sqlite3.connect('research_platform.db')
            cursor = conn.cursor()
            cursor.execute('SELECT setting_value FROM system_settings WHERE setting_key = ?', ('monitoring_interval',))
            result = cursor.fetchone()
            conn.close()

            if result:
                return int(result[0])
            else:
                return 1800  # 默认30分钟
        except Exception as e:
            logger.error(f"获取监控间隔设置失败: {e}")
            return 1800

    def update_monitoring_interval(self, interval_seconds):
        """更新监控间隔设置"""
        try:
            conn = sqlite3.connect('research_platform.db')
            cursor = conn.cursor()
            cursor.execute('''
                UPDATE system_settings
                SET setting_value = ?, updated_at = CURRENT_TIMESTAMP
                WHERE setting_key = ?
            ''', (str(interval_seconds), 'monitoring_interval'))
            conn.commit()
            conn.close()

            self.current_interval = interval_seconds
            logger.info(f"✅ 监控间隔已更新为 {interval_seconds} 秒")
            return True
        except Exception as e:
            logger.error(f"更新监控间隔设置失败: {e}")
            return False

    def start_monitoring(self):
        """启动监控服务"""
        if not self.running:
            self.running = True
            self.thread = threading.Thread(target=self._monitoring_loop, daemon=True)
            self.thread.start()
            logger.info(f"🚀 监控服务已启动 - 支持大规模监控，检查间隔: {self.current_interval}秒")

    def _monitoring_loop(self):
        """监控循环 - 使用可配置的时间间隔"""
        while self.running:
            try:
                self._check_researchers_batch()
                # 使用当前设置的间隔时间
                time.sleep(self.current_interval)
            except Exception as e:
                logger.error(f"监控循环错误: {e}")
                time.sleep(60)  # 出错时等待1分钟后重试

    def _check_researchers_batch(self):
        """批量检查正在监控的研究者"""
        conn = sqlite3.connect('research_platform.db')
        cursor = conn.cursor()

        cursor.execute('SELECT id, name, x_account FROM researchers WHERE is_monitoring = 1')
        researchers = cursor.fetchall()
        conn.close()

        logger.info(f"🔍 开始检查 {len(researchers)} 位研究者的内容")

        # 分批处理，避免同时处理过多研究者
        batch_size = 50  # 每批处理50个
        for i in range(0, len(researchers), batch_size):
            batch = researchers[i:i + batch_size]
            self._process_researcher_batch(batch)
            time.sleep(5)  # 批次间休息5秒

    def _process_researcher_batch(self, researchers_batch):
        """处理一批研究者"""
        for researcher_id, name, x_account in researchers_batch:
            try:
                tweets = twitter_api.get_user_tweets(x_account, max_results=5)

                if not tweets:
                    continue

                conn = sqlite3.connect('research_platform.db')
                cursor = conn.cursor()

                new_tweets_count = 0
                for tweet in tweets:
                    cursor.execute('''
                        INSERT OR IGNORE INTO x_content
                        (researcher_id, tweet_id, content, likes_count, retweets_count, replies_count, created_at)
                        VALUES (?, ?, ?, ?, ?, ?, ?)
                    ''', (
                        researcher_id, tweet['id'], tweet['content'],
                        tweet['likes'], tweet['retweets'], tweet['replies'],
                        tweet['created_at']
                    ))
                    if cursor.rowcount > 0:
                        new_tweets_count += 1

                # 更新最后检查时间
                cursor.execute('''
                    UPDATE monitoring_tasks SET last_check = CURRENT_TIMESTAMP
                    WHERE researcher_id = ?
                ''', (researcher_id,))

                conn.commit()
                conn.close()

                if new_tweets_count > 0:
                    logger.info(f"✅ {name} 更新了 {new_tweets_count} 条新内容")

            except Exception as e:
                logger.error(f"检查 {name} 时出错: {e}")
                time.sleep(1)  # 出错时稍作等待

# 初始化
try:
    researcher_manager = ResearcherManager()
    logger.info("✅ 研究者管理器初始化成功")
except Exception as e:
    logger.error(f"❌ 研究者管理器初始化失败: {e}")
    researcher_manager = None

try:
    twitter_api = TwitterAPI()
    logger.info("✅ Twitter API初始化完成")
except Exception as e:
    logger.error(f"❌ Twitter API初始化失败: {e}")
    twitter_api = None

# 初始化监控服务
try:
    monitoring_service = MonitoringService()
    logger.info("✅ 监控服务初始化成功")
except Exception as e:
    logger.error(f"❌ 监控服务初始化失败: {e}")
    monitoring_service = None

# API路由
@app.route('/')
def index():
    """主页路由 - 返回HTML模板"""
    try:
        return render_template('index.html')
    except Exception as e:
        logger.error(f"主页加载失败: {e}")
        return f"模板加载失败: {str(e)}<br>请确保 templates/index.html 文件存在", 500

@app.route('/api/researchers')
def get_researchers():
    """获取研究者列表 - 支持分页处理大量数据"""
    if not researcher_manager:
        return jsonify({'error': 'System not properly initialized'}), 500

    conn = sqlite3.connect('research_platform.db')
    cursor = conn.cursor()

    search_query = request.args.get('search', '')
    page = request.args.get('page', 1, type=int)
    per_page = request.args.get('per_page', 50, type=int)

    # 限制每页最大数量
    per_page = min(per_page, 200)
    offset = (page - 1) * per_page

    try:
        if search_query:
            # 搜索查询
            count_query = '''
                SELECT COUNT(*) FROM researchers
                WHERE name LIKE ? OR company LIKE ? OR research_focus LIKE ?
            '''
            cursor.execute(count_query, (f'%{search_query}%', f'%{search_query}%', f'%{search_query}%'))
            total_count = cursor.fetchone()[0]

            data_query = '''
                SELECT * FROM researchers
                WHERE name LIKE ? OR company LIKE ? OR research_focus LIKE ?
                ORDER BY rank LIMIT ? OFFSET ?
            '''
            cursor.execute(data_query, (f'%{search_query}%', f'%{search_query}%', f'%{search_query}%', per_page, offset))
        else:
            # 普通查询
            cursor.execute('SELECT COUNT(*) FROM researchers')
            total_count = cursor.fetchone()[0]

            cursor.execute('SELECT * FROM researchers ORDER BY rank LIMIT ? OFFSET ?', (per_page, offset))

        researchers = []
        for row in cursor.fetchall():
            researchers.append({
                'id': row[0], 'rank': row[1], 'name': row[2], 'country': row[3],
                'company': row[4], 'research_focus': row[5], 'x_account': row[6],
                'followers_count': row[7], 'following_count': row[8],
                'is_monitoring': bool(row[10]), 'is_special_focus': bool(row[11])
            })

        conn.close()

        return jsonify({
            'researchers': researchers,
            'pagination': {
                'page': page,
                'per_page': per_page,
                'total': total_count,
                'pages': (total_count + per_page - 1) // per_page
            }
        })

    except Exception as e:
        logger.error(f"获取研究者列表失败: {e}")
        conn.close()
        return jsonify({'error': str(e)}), 500

@app.route('/api/researcher/<int:researcher_id>')
def get_researcher_detail(researcher_id):
    """获取研究者详情"""
    conn = sqlite3.connect('research_platform.db')
    cursor = conn.cursor()

    cursor.execute('SELECT * FROM researchers WHERE id = ?', (researcher_id,))
    researcher_row = cursor.fetchone()

    if not researcher_row:
        conn.close()
        return jsonify({'error': 'Researcher not found'}), 404

    researcher = {
        'id': researcher_row[0], 'rank': researcher_row[1], 'name': researcher_row[2],
        'country': researcher_row[3], 'company': researcher_row[4],
        'research_focus': researcher_row[5], 'x_account': researcher_row[6],
        'followers_count': researcher_row[7], 'following_count': researcher_row[8],
        'is_monitoring': bool(researcher_row[10]), 'is_special_focus': bool(researcher_row[11])
    }

    # 获取最新内容
    cursor.execute('''
        SELECT * FROM x_content WHERE researcher_id = ?
        ORDER BY created_at DESC LIMIT 10
    ''', (researcher_id,))
    content_rows = cursor.fetchall()

    conn.close()

    recent_content = [
        {
            'id': c[0], 'content': c[3], 'likes': c[5],
            'retweets': c[6], 'replies': c[7], 'created_at': c[8]
        } for c in content_rows
    ]

    return jsonify({
        'researcher': researcher,
        'recent_content': recent_content
    })

@app.route('/api/researcher/<int:researcher_id>', methods=['DELETE'])
def delete_researcher(researcher_id):
    """删除指定的研究者及其所有相关数据"""
    try:
        conn = sqlite3.connect('research_platform.db')
        cursor = conn.cursor()

        cursor.execute("PRAGMA foreign_keys = ON;")
        cursor.execute('DELETE FROM researchers WHERE id = ?', (researcher_id,))

        conn.commit()

        if cursor.rowcount > 0:
            logger.info(f"✅ 成功删除研究者 ID: {researcher_id}")
            return jsonify({'message': f'成功删除研究者 ID: {researcher_id}'}), 200
        else:
            logger.warning(f"⚠️ 尝试删除一个不存在的研究者 ID: {researcher_id}")
            return jsonify({'error': 'Researcher not found'}), 404

    except Exception as e:
        logger.error(f"❌ 删除研究者 {researcher_id} 时发生错误: {e}")
        return jsonify({'error': 'Internal server error'}), 500
    finally:
        if conn:
            conn.close()

@app.route('/api/content')
def get_content():
    """获取所有内容 - 支持分页"""
    conn = sqlite3.connect('research_platform.db')
    cursor = conn.cursor()

    page = request.args.get('page', 1, type=int)
    per_page = request.args.get('per_page', 20, type=int)
    per_page = min(per_page, 100)  # 限制最大每页数量
    offset = (page - 1) * per_page

    try:
        # 获取总数
        cursor.execute('SELECT COUNT(*) FROM x_content')
        total_count = cursor.fetchone()[0]

        query = '''
            SELECT c.id, c.content, c.content_type, c.likes_count, c.retweets_count,
                   c.replies_count, c.created_at, c.collected_at, r.name, r.x_account
            FROM x_content c
            JOIN researchers r ON c.researcher_id = r.id
            ORDER BY c.created_at DESC
            LIMIT ? OFFSET ?
        '''

        cursor.execute(query, (per_page, offset))
        content_list = []

        for row in cursor.fetchall():
            content_list.append({
                'id': row[0],
                'content': row[1],
                'content_type': row[2],
                'likes_count': row[3],
                'retweets_count': row[4],
                'replies_count': row[5],
                'created_at': row[6],
                'collected_at': row[7],
                'author_name': row[8] or 'Unknown',
                'author_handle': row[9] or '@unknown'
            })

        conn.close()

        # 如果是简单请求（无分页参数），返回简单格式
        if page == 1 and per_page == 20:
            return jsonify(content_list)

        return jsonify({
            'content': content_list,
            'pagination': {
                'page': page,
                'per_page': per_page,
                'total': total_count,
                'pages': (total_count + per_page - 1) // per_page
            }
        })

    except Exception as e:
        logger.error(f"获取内容列表失败: {e}")
        conn.close()
        return jsonify({'error': str(e)}), 500

@app.route('/api/start_monitoring', methods=['POST'])
def start_monitoring_route():
    """开始监控指定研究者 - 支持批量操作"""
    data = request.get_json()
    researcher_ids = data.get('researcher_ids', [])

    if not researcher_ids:
        return jsonify({'error': 'No researchers selected'}), 400

    if len(researcher_ids) > 1000:  # 单次最多1000个
        return jsonify({'error': 'Too many researchers selected at once (max: 1000)'}), 400

    conn = sqlite3.connect('research_platform.db')
    cursor = conn.cursor()

    success_count = 0
    failed_ids = []

    # 开始事务
    cursor.execute('BEGIN TRANSACTION')

    try:
        for researcher_id in researcher_ids:
            try:
                # 更新研究者监控状态
                cursor.execute('UPDATE researchers SET is_monitoring = 1, updated_at = CURRENT_TIMESTAMP WHERE id = ?', (researcher_id,))

                # 创建监控任务
                cursor.execute('INSERT OR REPLACE INTO monitoring_tasks (researcher_id, status, last_check) VALUES (?, \'active\', CURRENT_TIMESTAMP)', (researcher_id,))

                success_count += 1

            except Exception as e:
                logger.error(f"启动监控研究者 {researcher_id} 失败: {e}")
                failed_ids.append(researcher_id)

        cursor.execute('COMMIT')

    except Exception as e:
        cursor.execute('ROLLBACK')
        logger.error(f"批量启动监控失败: {e}")
        return jsonify({'error': 'Failed to start monitoring'}), 500

    finally:
        conn.close()

    # 确保监控服务正在运行
    if monitoring_service:
        monitoring_service.start_monitoring()

    response_data = {
        'message': f'成功启动监控 {success_count} 位研究者',
        'monitoring_count': success_count
    }

    if failed_ids:
        response_data['failed_ids'] = failed_ids
        response_data['message'] += f', {len(failed_ids)} 位失败'

    return jsonify(response_data)

@app.route('/api/stop_monitoring', methods=['POST'])
def stop_monitoring_route():
    """停止监控指定研究者"""
    data = request.get_json()
    researcher_ids = data.get('researcher_ids', [])

    if len(researcher_ids) > 1000:
        return jsonify({'error': 'Too many researchers selected at once (max: 1000)'}), 400

    conn = sqlite3.connect('research_platform.db')
    cursor = conn.cursor()

    cursor.execute('BEGIN TRANSACTION')

    try:
        for researcher_id in researcher_ids:
            cursor.execute('UPDATE researchers SET is_monitoring = 0, updated_at = CURRENT_TIMESTAMP WHERE id = ?', (researcher_id,))
            cursor.execute('UPDATE monitoring_tasks SET status = \'inactive\' WHERE researcher_id = ?', (researcher_id,))

        cursor.execute('COMMIT')

    except Exception as e:
        cursor.execute('ROLLBACK')
        logger.error(f"批量停止监控失败: {e}")
        return jsonify({'error': 'Failed to stop monitoring'}), 500

    finally:
        conn.close()

    return jsonify({'message': f'已停止监控 {len(researcher_ids)} 位研究者'})

@app.route('/api/fetch_content/<int:researcher_id>', methods=['POST'])
def fetch_researcher_content(researcher_id):
    """立即获取指定研究者的最新内容"""
    conn = sqlite3.connect('research_platform.db')
    cursor = conn.cursor()

    cursor.execute('SELECT name, x_account FROM researchers WHERE id = ?', (researcher_id,))
    researcher = cursor.fetchone()

    if not researcher:
        conn.close()
        return jsonify({'error': 'Researcher not found'}), 404

    name, x_account = researcher

    try:
        # 获取最新推文
        tweets = twitter_api.get_user_tweets(x_account, max_results=10) if twitter_api else []

        new_content_count = 0
        if tweets:
            for tweet in tweets:
                cursor.execute('''
                    INSERT OR IGNORE INTO x_content
                    (researcher_id, tweet_id, content, likes_count, retweets_count, replies_count, created_at)
                    VALUES (?, ?, ?, ?, ?, ?, ?)
                ''', (
                    researcher_id, tweet['id'], tweet['content'],
                    tweet['likes'], tweet['retweets'], tweet['replies'],
                    tweet['created_at']
                ))

                if cursor.rowcount > 0:
                    new_content_count += 1

            conn.commit()

        conn.close()

        message = f'成功获取 {name} 的内容。' if tweets else f'未找到 {name} 的新内容。'
        return jsonify({
            'message': message,
            'new_content_count': new_content_count,
            'total_fetched': len(tweets)
        })

    except Exception as e:
        conn.close()
        logger.error(f"获取 {name} 内容失败: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/monitoring_settings')
def get_monitoring_settings():
    """获取监控设置"""
    try:
        conn = sqlite3.connect('research_platform.db')
        cursor = conn.cursor()

        cursor.execute('SELECT setting_key, setting_value, description FROM system_settings')
        settings = {}
        for row in cursor.fetchall():
            settings[row[0]] = {
                'value': row[1],
                'description': row[2]
            }

        conn.close()

        # 计算当前间隔的人性化显示
        interval_seconds = int(settings.get('monitoring_interval', {}).get('value', 1800))
        interval_display = format_interval(interval_seconds)

        return jsonify({
            'monitoring_interval': interval_seconds,
            'interval_display': interval_display,
            'settings': settings,
            'predefined_intervals': [
                {'value': 1800, 'label': '30分钟', 'description': '高频监控，适合热点关注'},
                {'value': 3600, 'label': '1小时', 'description': '标准监控，平衡效率与时效'},
                {'value': 7200, 'label': '2小时', 'description': '中等频率，节省资源'},
                {'value': 21600, 'label': '6小时', 'description': '低频监控，适合长期观察'},
                {'value': 43200, 'label': '12小时', 'description': '每日两次检查'},
                {'value': 86400, 'label': '24小时', 'description': '每日一次，最节省资源'}
            ]
        })

    except Exception as e:
        logger.error(f"获取监控设置失败: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/monitoring_settings', methods=['POST'])
def update_monitoring_settings():
    """更新监控设置"""
    try:
        data = request.get_json()
        interval_seconds = data.get('monitoring_interval')

        if not interval_seconds or not isinstance(interval_seconds, int):
            return jsonify({'error': '无效的监控间隔值'}), 400

        if interval_seconds < 300:  # 最小5分钟
            return jsonify({'error': '监控间隔不能少于5分钟（300秒）'}), 400

        if interval_seconds > 604800:  # 最大7天
            return jsonify({'error': '监控间隔不能超过7天（604800秒）'}), 400

        # 更新监控服务的间隔
        if monitoring_service and monitoring_service.update_monitoring_interval(interval_seconds):
            return jsonify({
                'message': f'监控间隔已更新为 {format_interval(interval_seconds)}',
                'new_interval': interval_seconds,
                'interval_display': format_interval(interval_seconds)
            })
        else:
            return jsonify({'error': '更新监控间隔失败'}), 500

    except Exception as e:
        logger.error(f"更新监控设置失败: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/analytics')
def get_analytics():
    """获取平台分析数据"""
    conn = sqlite3.connect('research_platform.db')
    cursor = conn.cursor()

    # 基础统计
    cursor.execute('SELECT COUNT(*) FROM researchers')
    total_researchers = cursor.fetchone()[0]

    cursor.execute('SELECT COUNT(*) FROM researchers WHERE is_monitoring = 1')
    monitoring_researchers = cursor.fetchone()[0]

    cursor.execute('SELECT COUNT(*) FROM x_content')
    total_content = cursor.fetchone()[0]

    cursor.execute('SELECT SUM(likes_count + retweets_count + replies_count) FROM x_content')
    total_engagement = cursor.fetchone()[0] or 0

    # 国家分布
    cursor.execute('SELECT country, COUNT(*) FROM researchers GROUP BY country')
    country_distribution = {k: v for k, v in cursor.fetchall() if k}

    # 公司分布
    cursor.execute('SELECT company, COUNT(*) FROM researchers GROUP BY company')
    company_distribution = {k: v for k, v in cursor.fetchall() if k}

    # 最近7天的内容趋势
    cursor.execute('''
        SELECT DATE(created_at), COUNT(*)
        FROM x_content
        WHERE created_at >= date('now', '-7 days')
        GROUP BY DATE(created_at)
        ORDER BY DATE(created_at)
    ''')
    content_trend = dict(cursor.fetchall())

    # 监控能力状态
    cursor.execute('SELECT MAX(rank) FROM researchers')
    max_capacity = 5000  # 最大支持容量
    current_capacity = cursor.fetchone()[0] or 0

    # 获取监控间隔设置
    monitoring_interval = 1800  # 默认值
    interval_display = "30分钟"
    try:
        cursor.execute('SELECT setting_value FROM system_settings WHERE setting_key = ?', ('monitoring_interval',))
        result = cursor.fetchone()
        if result:
            monitoring_interval = int(result[0])
            interval_display = format_interval(monitoring_interval)
    except Exception as e:
        logger.error(f"获取监控间隔设置失败: {e}")

    conn.close()

    return jsonify({
        'total_researchers': total_researchers,
        'monitoring_researchers': monitoring_researchers,
        'total_content': total_content,
        'total_engagement': total_engagement,
        'country_distribution': country_distribution,
        'company_distribution': company_distribution,
        'content_trend': content_trend,
        'api_status': 'connected' if twitter_api and twitter_api.client else 'disconnected',
        'monitoring_active': monitoring_service and monitoring_service.running,
        'monitoring_interval': monitoring_interval,
        'interval_display': interval_display,
        'capacity': {
            'current': total_researchers,
            'monitoring': monitoring_researchers,
            'max_supported': max_capacity,
            'utilization': f"{(total_researchers/max_capacity)*100:.1f}%"
        }
    })

@app.route('/api/upload_excel', methods=['POST'])
def upload_excel():
    """上传Excel文件 - 增强错误处理和批量导入"""
    if 'file' not in request.files:
        return jsonify({'error': 'No file uploaded'}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400

    # 检查文件类型
    if not file.filename.lower().endswith(('.xlsx', '.xls')):
        return jsonify({'error': 'Please upload an Excel file (.xlsx or .xls)'}), 400

    try:
        import openpyxl
        workbook = openpyxl.load_workbook(file)
        worksheet = workbook.active

        logger.info(f"📊 开始处理Excel文件，共 {worksheet.max_row - 1} 行数据")

        conn = sqlite3.connect('research_platform.db')
        cursor = conn.cursor()

        # 开始事务
        cursor.execute('BEGIN TRANSACTION')

        added_count = 0
        error_count = 0
        skipped_count = 0
        error_details = []

        # 批量处理数据
        batch_size = 100
        batch_data = []

        for row_num, row in enumerate(worksheet.iter_rows(min_row=2, values_only=True), start=2):
            try:
                # 数据验证
                if not row or len(row) < 6:
                    skipped_count += 1
                    logger.warning(f"第 {row_num} 行：数据不完整，跳过")
                    continue

                if not row[1]:  # 名字不能为空
                    skipped_count += 1
                    logger.warning(f"第 {row_num} 行：研究者姓名为空，跳过")
                    continue

                # 清理和验证数据
                rank = row[0] if row[0] is not None else row_num - 1
                name = str(row[1]).strip() if row[1] else ''
                country = str(row[2]).strip() if row[2] else ''
                company = str(row[3]).strip() if row[3] else ''
                research_focus = str(row[4]).strip() if row[4] else ''
                x_account = str(row[5]).strip() if row[5] else ''

                # 确保 X 账号格式正确
                if x_account and not x_account.startswith('@'):
                    x_account = '@' + x_account

                batch_data.append((rank, name, country, company, research_focus, x_account))

                # 达到批量大小时执行插入
                if len(batch_data) >= batch_size:
                    added_count += insert_researcher_batch(cursor, batch_data, error_details)
                    batch_data = []

            except Exception as e:
                error_count += 1
                error_msg = f"第 {row_num} 行处理失败: {str(e)}"
                logger.error(error_msg)
                error_details.append(error_msg)

                if error_count > 50:  # 如果错误太多，停止处理
                    logger.error("错误过多，停止处理文件")
                    break

        # 处理剩余的批量数据
        if batch_data:
            added_count += insert_researcher_batch(cursor, batch_data, error_details)

        # 提交事务
        cursor.execute('COMMIT')
        conn.close()

        total_processed = worksheet.max_row - 1

        logger.info(f"✅ Excel导入完成: 成功 {added_count}, 跳过 {skipped_count}, 错误 {error_count}")

        response_data = {
            'message': f'Excel文件处理完成',
            'total_rows': total_processed,
            'imported': added_count,
            'skipped': skipped_count,
            'errors': error_count
        }

        if error_details and len(error_details) <= 20:  # 只返回前20个错误
            response_data['error_details'] = error_details[:20]

        return jsonify(response_data)

    except Exception as e:
        logger.error(f"❌ Excel文件处理失败: {e}")
        return jsonify({
            'error': f'文件处理失败: {str(e)}',
            'suggestion': '请检查文件格式，确保包含必要的列：排名、姓名、国家、公司、研究领域、X账号'
        }), 500

@app.route('/api/special_focus', methods=['POST'])
def set_special_focus():
    """设置特别关注"""
    try:
        data = request.get_json()
        researcher_ids = data.get('researcher_ids', [])
        is_special = data.get('is_special', True)

        if not researcher_ids:
            return jsonify({'error': 'No researchers selected'}), 400

        conn = sqlite3.connect('research_platform.db')
        cursor = conn.cursor()

        success_count = 0
        for researcher_id in researcher_ids:
            cursor.execute('''
                UPDATE researchers
                SET is_special_focus = ?, updated_at = CURRENT_TIMESTAMP
                WHERE id = ?
            ''', (is_special, researcher_id))
            success_count += 1

        conn.commit()
        conn.close()

        action = "设为特别关注" if is_special else "取消特别关注"
        return jsonify({
            'message': f'成功{action} {success_count} 位研究者',
            'success_count': success_count
        })

    except Exception as e:
        logger.error(f"设置特别关注失败: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/special_focus')
def get_special_focus():
    """获取特别关注的研究者列表"""
    try:
        conn = sqlite3.connect('research_platform.db')
        cursor = conn.cursor()

        cursor.execute('''
            SELECT id, rank, name, country, company, research_focus, x_account,
                   followers_count, following_count, avatar_url, is_monitoring, is_special_focus
            FROM researchers
            WHERE is_special_focus = 1
            ORDER BY name
        ''')

        researchers = []
        for row in cursor.fetchall():
            researchers.append({
                'id': row[0], 'rank': row[1], 'name': row[2], 'country': row[3],
                'company': row[4], 'research_focus': row[5], 'x_account': row[6],
                'followers_count': row[7], 'following_count': row[8],
                'avatar_url': row[9], 'is_monitoring': bool(row[10]),
                'is_special_focus': bool(row[11])
            })

        conn.close()
        logger.info(f"获取特别关注列表成功，共 {len(researchers)} 位")
        return jsonify(researchers)

    except Exception as e:
        logger.error(f"获取特别关注列表失败: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/update_user_info/<int:researcher_id>', methods=['POST'])
def update_user_info(researcher_id):
    """更新研究者的用户信息（关注者等数据）"""
    try:
        conn = sqlite3.connect('research_platform.db')
        cursor = conn.cursor()

        cursor.execute('SELECT name, x_account FROM researchers WHERE id = ?', (researcher_id,))
        researcher = cursor.fetchone()

        if not researcher:
            conn.close()
            return jsonify({'error': 'Researcher not found'}), 404

        name, x_account = researcher

        # 获取用户信息
        user_info = twitter_api.get_user_info(x_account) if twitter_api else None

        if user_info:
            # 更新数据库中的用户信息 - 直接存储数字而不是格式化字符串
            cursor.execute('''
                UPDATE researchers
                SET followers_count = ?, following_count = ?, updated_at = CURRENT_TIMESTAMP
                WHERE id = ?
            ''', (str(user_info['followers_count']), str(user_info['following_count']), researcher_id))

            conn.commit()
            conn.close()

            return jsonify({
                'message': f'成功更新 {name} 的用户信息',
                'user_info': user_info
            })
        else:
            conn.close()
            return jsonify({
                'message': f'无法获取 {name} 的用户信息（可能是API限制或网络问题）',
                'user_info': None
            })

    except Exception as e:
        logger.error(f"更新用户信息失败: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/update_all_user_info', methods=['POST'])
def update_all_user_info():
    """批量更新所有研究者的用户信息"""
    try:
        conn = sqlite3.connect('research_platform.db')
        cursor = conn.cursor()

        cursor.execute('SELECT id, name, x_account FROM researchers ORDER BY id')
        researchers = cursor.fetchall()

        updated_count = 0
        failed_count = 0

        for researcher_id, name, x_account in researchers:
            try:
                # 获取用户信息
                user_info = twitter_api.get_user_info(x_account) if twitter_api else None

                if user_info:
                    # 更新数据库
                    cursor.execute('''
                        UPDATE researchers
                        SET followers_count = ?, following_count = ?, updated_at = CURRENT_TIMESTAMP
                        WHERE id = ?
                    ''', (str(user_info['followers_count']), str(user_info['following_count']), researcher_id))

                    updated_count += 1
                    logger.info(f"✅ 更新 {name}: {user_info['followers_count']} 关注者, {user_info['following_count']} 正在关注")
                else:
                    failed_count += 1
                    logger.warning(f"⚠️ 无法获取 {name} 的用户信息")

                # 添加延迟，避免API限制
                time.sleep(2)  # 增加延迟到2秒

            except Exception as e:
                failed_count += 1
                logger.error(f"❌ 更新 {name} 失败: {e}")

        conn.commit()
        conn.close()

        return jsonify({
            'message': f'批量更新完成: 成功 {updated_count} 个，失败 {failed_count} 个',
            'updated_count': updated_count,
            'failed_count': failed_count
        })

    except Exception as e:
        logger.error(f"批量更新用户信息失败: {e}")
        return jsonify({'error': str(e)}), 500

# 在app.py中添加一个新的调试端点
@app.route('/api/debug_fetch_historical/<int:researcher_id>', methods=['POST'])
def debug_fetch_historical_content(researcher_id):
    """调试版本的历史内容抓取 - 输出详细调试信息"""
    try:
        data = request.get_json()
        start_date = data.get('start_date')
        end_date = data.get('end_date', datetime.now().strftime('%Y-%m-%d'))
        max_results = data.get('max_results', 50)

        debug_log = []  # 收集调试信息

        conn = sqlite3.connect('research_platform.db')
        cursor = conn.cursor()

        cursor.execute('SELECT name, x_account FROM researchers WHERE id = ?', (researcher_id,))
        researcher = cursor.fetchone()

        if not researcher:
            conn.close()
            return jsonify({'error': 'Researcher not found'}), 404

        name, x_account = researcher
        debug_log.append(f"📋 研究者: {name} ({x_account})")

        # 时间范围处理
        if start_date:
            start_time = datetime.strptime(start_date, '%Y-%m-%d').replace(tzinfo=timezone.utc)
        else:
            start_time = datetime.now(timezone.utc) - timedelta(days=7)

        end_time = datetime.strptime(end_date, '%Y-%m-%d').replace(tzinfo=timezone.utc)

        debug_log.append(f"⏰ 时间范围: {start_time.strftime('%Y-%m-%d %H:%M:%S UTC')} 到 {end_time.strftime('%Y-%m-%d %H:%M:%S UTC')}")
        debug_log.append(f"📊 时间跨度: {(end_time - start_time).days} 天")
        debug_log.append(f"🎯 最大结果数: {max_results}")

        # 检查Twitter API状态
        if not twitter_api:
            debug_log.append("❌ Twitter API未初始化")
            return jsonify({
                'error': 'Twitter API未初始化',
                'debug_log': debug_log
            }), 500

        debug_log.append(f"🔗 Twitter API状态: {'正常' if twitter_api.api_working else '异常'}")
        debug_log.append(f"🧪 连接已测试: {'是' if twitter_api.connection_tested else '否'}")

        # 先测试用户信息获取
        debug_log.append("👤 开始获取用户信息...")
        user_info = twitter_api.get_user_info(x_account)
        if user_info:
            debug_log.append(f"✅ 用户信息获取成功: @{user_info['username']}")
            debug_log.append(f"   📊 关注者: {user_info['followers_count']:,}")
            debug_log.append(f"   📊 正在关注: {user_info['following_count']:,}")
            debug_log.append(f"   📊 推文总数: {user_info['tweet_count']:,}")
            debug_log.append(f"   🔒 认证状态: {'已认证' if user_info.get('verified') else '未认证'}")
        else:
            debug_log.append("❌ 用户信息获取失败 - 可能是私人账户或不存在")

        # 调用推文获取方法
        debug_log.append("🐦 开始调用推文获取方法...")
        tweets = twitter_api.get_user_tweets(
            x_account,
            max_results=max_results,
            start_time=start_time,
            end_time=end_time
        )

        debug_log.append(f"📥 推文获取完成，返回 {len(tweets)} 条")

        # 如果没有获取到推文，尝试不同的参数
        if not tweets:
            debug_log.append("⚠️ 未获取到推文，尝试扩大时间范围...")

            # 尝试更大的时间范围
            extended_start_time = datetime.now(timezone.utc) - timedelta(days=90)
            debug_log.append(f"🔄 扩大时间范围到90天: {extended_start_time.strftime('%Y-%m-%d')} 到 {end_time.strftime('%Y-%m-%d')}")

            extended_tweets = twitter_api.get_user_tweets(
                x_account,
                max_results=max_results,
                start_time=extended_start_time,
                end_time=end_time
            )

            debug_log.append(f"📥 扩大时间范围后获取: {len(extended_tweets)} 条")

            if extended_tweets:
                tweets = extended_tweets
                debug_log.append("✅ 扩大时间范围成功获取到推文")
            else:
                debug_log.append("⚠️ 即使扩大时间范围也未获取到推文")

                # 再尝试一次，不设置时间限制
                debug_log.append("🔄 尝试不设置时间限制...")
                no_time_tweets = twitter_api.get_user_tweets(
                    x_account,
                    max_results=max_results
                    # 不设置start_time和end_time
                )

                debug_log.append(f"📥 不设置时间限制获取: {len(no_time_tweets)} 条")

                if no_time_tweets:
                    tweets = no_time_tweets
                    debug_log.append("✅ 不设置时间限制成功获取到推文")

        # 分析推文内容
        if tweets:
            debug_log.append("🔍 分析获取到的推文:")

            # 按类型统计
            type_stats = {}
            date_range = {'earliest': None, 'latest': None}

            for tweet in tweets:
                tweet_type = tweet.get('type', 'unknown')
                type_stats[tweet_type] = type_stats.get(tweet_type, 0) + 1

                if tweet.get('created_at'):
                    tweet_date = datetime.fromisoformat(tweet['created_at'].replace('Z', '+00:00'))
                    if not date_range['earliest'] or tweet_date < date_range['earliest']:
                        date_range['earliest'] = tweet_date
                    if not date_range['latest'] or tweet_date > date_range['latest']:
                        date_range['latest'] = tweet_date

            debug_log.append(f"   📊 推文类型统计: {type_stats}")
            if date_range['earliest'] and date_range['latest']:
                debug_log.append(f"   📅 推文时间范围: {date_range['earliest'].strftime('%Y-%m-%d')} 到 {date_range['latest'].strftime('%Y-%m-%d')}")

            # 显示前几条推文的详细信息
            debug_log.append("   📝 推文样本:")
            for i, tweet in enumerate(tweets[:3]):
                debug_log.append(f"      {i+1}. [{tweet.get('type', 'unknown')}] {tweet.get('created_at', 'N/A')[:19]} | 👍{tweet.get('likes', 0)} | {tweet.get('content', '')[:60]}...")

        # 数据库保存
        new_content_count = 0
        duplicate_count = 0
        error_count = 0

        if tweets:
            debug_log.append("💾 开始保存到数据库...")

            for i, tweet in enumerate(tweets):
                try:
                    # 检查是否已存在
                    cursor.execute('SELECT id FROM x_content WHERE tweet_id = ?', (tweet['id'],))
                    existing = cursor.fetchone()

                    if existing:
                        duplicate_count += 1
                        continue

                    # 存储媒体URL
                    media_urls_json = json.dumps(tweet.get('media_urls', []))

                    # 插入新记录
                    cursor.execute('''
                        INSERT INTO x_content
                        (researcher_id, tweet_id, content, content_type, likes_count, retweets_count,
                         replies_count, created_at, is_historical, media_urls)
                        VALUES (?, ?, ?, ?, ?, ?, ?, ?, 1, ?)
                    ''', (
                        researcher_id, tweet['id'], tweet['content'], tweet.get('type', 'text'),
                        tweet['likes'], tweet['retweets'], tweet['replies'],
                        tweet['created_at'], media_urls_json
                    ))

                    new_content_count += 1

                except Exception as e:
                    error_count += 1
                    debug_log.append(f"❌ 保存推文 {i+1} 失败: {str(e)}")

            conn.commit()
            debug_log.append(f"💾 数据库保存完成: 新增{new_content_count}条，重复{duplicate_count}条，错误{error_count}条")

        conn.close()

        # 生成结果消息
        if tweets:
            result_message = f'✅ {name} 调试抓取完成: 获取{len(tweets)}条，新增{new_content_count}条，重复{duplicate_count}条'
        else:
            result_message = f'⚠️ {name} 调试抓取完成但未获取到推文'

            # 分析可能的原因
            debug_log.append("🔍 未获取到推文的可能原因分析:")
            if not user_info:
                debug_log.append("   1. 用户账户不存在或无法访问（私人账户）")
            elif user_info.get('tweet_count', 0) == 0:
                debug_log.append("   2. 用户的推文总数为0（新账户或删除了所有推文）")
            else:
                debug_log.append("   3. 用户在指定时间范围内没有发布推文")
                debug_log.append("   4. 推文被API的某些过滤条件排除")
                debug_log.append("   5. API权限不足，无法访问该用户的推文")

        debug_log.append(f"🎉 调试完成: {result_message}")

        return jsonify({
            'message': result_message,
            'new_content_count': new_content_count,
            'total_fetched': len(tweets),
            'duplicate_count': duplicate_count,
            'error_count': error_count,
            'debug_log': debug_log,
            'user_info': user_info,
            'debug_details': {
                'researcher_name': name,
                'x_account': x_account,
                'time_range_days': (end_time - start_time).days,
                'api_working': twitter_api and twitter_api.api_working,
                'connection_tested': twitter_api and twitter_api.connection_tested,
                'user_accessible': user_info is not None,
                'user_tweet_count': user_info.get('tweet_count', 0) if user_info else 0
            }
        })

    except Exception as e:
        logger.error(f"💥 调试抓取历史内容时发生严重错误: {e}")
        return jsonify({
            'error': str(e),
            'debug_log': debug_log if 'debug_log' in locals() else [],
            'error_type': type(e).__name__
        }), 500

# 同时添加一个简单的测试不同用户的端点
@app.route('/api/test_user_tweets/<username>', methods=['POST'])
def test_user_tweets(username):
    """测试特定用户名的推文获取"""
    try:
        data = request.get_json() or {}
        max_results = data.get('max_results', 20)
        days = data.get('days', 7)

        if not twitter_api:
            return jsonify({'error': 'Twitter API未初始化'}), 500

        # 清理用户名
        clean_username = username.replace('@', '').strip()

        logger.info(f"🧪 测试用户推文获取: @{clean_username}")

        # 设置时间范围
        start_time = datetime.now(timezone.utc) - timedelta(days=days)
        end_time = datetime.now(timezone.utc)

        # 获取用户信息
        user_info = twitter_api.get_user_info(clean_username)

        # 获取推文
        tweets = twitter_api.get_user_tweets(
            clean_username,
            max_results=max_results,
            start_time=start_time,
            end_time=end_time
        )

        return jsonify({
            'username': clean_username,
            'user_info': user_info,
            'tweets_count': len(tweets),
            'tweets_sample': tweets[:3],  # 返回前3条作为样本
            'time_range': f"{start_time.strftime('%Y-%m-%d')} 到 {end_time.strftime('%Y-%m-%d')}",
            'success': len(tweets) > 0,
            'message': f"成功获取 @{clean_username} 的 {len(tweets)} 条推文" if tweets else f"未获取到 @{clean_username} 的推文"
        })

    except Exception as e:
        logger.error(f"测试用户推文获取失败: {e}")
        return jsonify({
            'error': str(e),
            'username': username
        }), 500

@app.route('/api/export_word/<int:researcher_id>')
def export_to_word(researcher_id):
    """导出研究者内容为Word文档"""
    try:
        from docx import Document
        from docx.shared import Inches
        import io

        conn = sqlite3.connect('research_platform.db')
        cursor = conn.cursor()

        # 获取研究者信息
        cursor.execute('SELECT name, x_account, research_focus FROM researchers WHERE id = ?', (researcher_id,))
        researcher = cursor.fetchone()

        if not researcher:
            conn.close()
            return jsonify({'error': 'Researcher not found'}), 404

        name, x_account, research_focus = researcher

        # 获取所有内容
        cursor.execute('''
            SELECT content, likes_count, retweets_count, replies_count,
                   created_at, media_urls
            FROM x_content
            WHERE researcher_id = ?
            ORDER BY created_at DESC
        ''', (researcher_id,))

        contents = cursor.fetchall()
        conn.close()

        # 创建Word文档
        doc = Document()

        # 添加标题
        title = doc.add_heading(f'{name} 内容记录', 0)

        # 添加基本信息
        doc.add_heading('基本信息', level=1)
        info_table = doc.add_table(rows=3, cols=2)
        info_table.style = 'Table Grid'

        info_table.cell(0, 0).text = '姓名'
        info_table.cell(0, 1).text = name
        info_table.cell(1, 0).text = 'X账号'
        info_table.cell(1, 1).text = x_account
        info_table.cell(2, 0).text = '研究领域'
        info_table.cell(2, 1).text = research_focus or '未知'

        # 添加内容
        doc.add_heading('内容记录', level=1)
        doc.add_paragraph(f'共收集 {len(contents)} 条内容，按时间倒序排列：')

        for i, content in enumerate(contents, 1):
            text, likes, retweets, replies, created_at, media_urls = content

            # 添加序号和时间
            heading = doc.add_heading(f'{i}. {created_at[:19] if created_at else "未知时间"}', level=2)

            # 添加内容
            doc.add_paragraph(text or '无文字内容')

            # 添加媒体信息
            if media_urls:
                try:
                    media_list = json.loads(media_urls)
                    if media_list:
                        doc.add_paragraph('包含媒体:')
                        for media in media_list:
                            media_type = media.get('type', 'unknown')
                            media_url = media.get('url', media.get('preview_url', ''))
                            if media_url:
                                doc.add_paragraph(f'• {media_type}: {media_url}', style='List Bullet')
                except:
                    pass

            # 添加互动数据
            stats_p = doc.add_paragraph()
            stats_p.add_run(f'👍 {likes} 点赞  🔄 {retweets} 转发  💬 {replies} 回复')

            # 添加分隔线
            if i < len(contents):
                doc.add_paragraph('─' * 50)

        # 保存到内存
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        filename = f"{name}_内容记录_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

        return send_file(
            file_stream,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except ImportError:
        return jsonify({
            'error': 'Word导出功能需要安装 python-docx 库',
            'solution': '请运行: pip install python-docx'
        }), 500
    except Exception as e:
        logger.error(f"导出Word失败: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/database_status')
def get_database_status():
    """获取数据库状态信息"""
    try:
        conn = sqlite3.connect('research_platform.db')
        cursor = conn.cursor()

        # 检查各表的记录数
        cursor.execute('SELECT COUNT(*) FROM researchers')
        researchers_count = cursor.fetchone()[0]

        cursor.execute('SELECT COUNT(*) FROM x_content')
        content_count = cursor.fetchone()[0]

        cursor.execute('SELECT COUNT(*) FROM monitoring_tasks')
        tasks_count = cursor.fetchone()[0]

        # 检查数据库文件大小
        import os
        db_size = os.path.getsize('research_platform.db') if os.path.exists('research_platform.db') else 0

        # 检查初始化状态
        cursor.execute('SELECT value FROM db_metadata WHERE key = ?', ('sample_data_loaded',))
        initialized = cursor.fetchone()

        conn.close()

        return jsonify({
            'database_file': 'research_platform.db',
            'file_exists': os.path.exists('research_platform.db'),
            'file_size_bytes': db_size,
            'file_size_mb': round(db_size / 1024 / 1024, 2),
            'tables': {
                'researchers': researchers_count,
                'x_content': content_count,
                'monitoring_tasks': tasks_count
            },
            'initialized': bool(initialized),
            'status': 'healthy' if researchers_count > 0 else 'empty'
        })

    except Exception as e:
        logger.error(f"获取数据库状态失败: {e}")
        return jsonify({'error': str(e), 'status': 'error'}), 500

@app.route('/api/system_status')
def get_system_status():
    """获取系统状态信息"""
    conn = sqlite3.connect('research_platform.db')
    cursor = conn.cursor()

    # 数据库统计
    cursor.execute('SELECT COUNT(*) FROM researchers')
    total_researchers = cursor.fetchone()[0]

    cursor.execute('SELECT COUNT(*) FROM researchers WHERE is_monitoring = 1')
    monitoring_count = cursor.fetchone()[0]

    cursor.execute('SELECT COUNT(*) FROM x_content')
    total_content = cursor.fetchone()[0]

    # 最近24小时的活动
    cursor.execute('''
        SELECT COUNT(*) FROM x_content
        WHERE collected_at >= datetime('now', '-1 day')
    ''')
    recent_content = cursor.fetchone()[0]

    # 获取监控间隔
    monitoring_interval = 1800
    try:
        cursor.execute('SELECT setting_value FROM system_settings WHERE setting_key = ?', ('monitoring_interval',))
        result = cursor.fetchone()
        if result:
            monitoring_interval = int(result[0])
    except Exception as e:
        logger.error(f"获取监控间隔失败: {e}")

    conn.close()

    return jsonify({
        'system_capacity': {
            'max_researchers': 5000,
            'current_researchers': total_researchers,
            'available_slots': 5000 - total_researchers,
            'utilization_percentage': (total_researchers / 5000) * 100
        },
        'monitoring_status': {
            'active_monitoring': monitoring_count,
            'max_concurrent': 1000,
            'service_running': monitoring_service and monitoring_service.running,
            'monitoring_interval': monitoring_interval,
            'interval_display': format_interval(monitoring_interval)
        },
        'data_statistics': {
            'total_content': total_content,
            'recent_24h': recent_content
        },
        'api_status': {
            'twitter_connected': twitter_api and twitter_api.client is not None,
            'twitter_working': twitter_api and twitter_api.api_working,
            'last_check': datetime.now().isoformat()
        }
    })

@app.route('/api/reset_sample_data', methods=['POST'])
def reset_sample_data():
    """重置示例数据（仅用于测试和恢复）"""
    try:
        if researcher_manager:
            conn = sqlite3.connect('research_platform.db')
            cursor = conn.cursor()

            # 删除现有示例数据（基于名字判断）
            sample_names = ['Ilya Sutskever', 'Noam Shazeer', 'Geoffrey Hinton', 'Alec Radford', 'Andrej Karpathy']
            for name in sample_names:
                cursor.execute('DELETE FROM researchers WHERE name = ?', (name,))

            # 重置初始化标记
            cursor.execute('DELETE FROM db_metadata WHERE key = ?', ('sample_data_loaded',))

            conn.commit()
            conn.close()

            # 重新加载示例数据
            researcher_manager.load_sample_data_if_empty()

            return jsonify({'message': '示例数据已重置', 'status': 'success'})
        else:
            return jsonify({'error': '研究者管理器未初始化'}), 500

    except Exception as e:
        logger.error(f"重置示例数据失败: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/health')
def health_check():
    """健康检查"""
    return jsonify({
        'status': 'healthy' if researcher_manager else 'partial',
        'timestamp': datetime.now().isoformat(),
        'twitter_api': 'connected' if twitter_api and twitter_api.client else 'disconnected',
        'twitter_working': twitter_api and twitter_api.api_working,
        'monitoring': 'active' if monitoring_service and monitoring_service.running else 'inactive',
        'capacity': '5000 researchers supported',
        'components': {
            'researcher_manager': 'ok' if researcher_manager else 'failed',
            'twitter_api': 'ok' if twitter_api else 'failed',
            'monitoring_service': 'ok' if monitoring_service else 'failed'
        }
    })

@app.route('/api/init_status')
def get_init_status():
    """获取初始化状态"""
    return jsonify({
        'initialized': bool(researcher_manager),
        'components': {
            'database': bool(researcher_manager),
            'twitter_api': bool(twitter_api),
            'monitoring': bool(monitoring_service)
        },
        'ready': bool(researcher_manager and twitter_api and monitoring_service)
    })

                            # 在app.py中添加这些详细的调试端点

# 安全版本的调试端点 - 添加到app.py的路由函数区域

@app.route('/api/safe_debug_tweets', methods=['POST'])
def safe_debug_tweets():
    """安全版本的推文调试 - 带超时和错误处理"""
    try:
        data = request.get_json() or {}
        username = data.get('username', '').replace('@', '').strip()
        max_results = min(data.get('max_results', 10), 20)  # 限制最大数量

        logger.info(f"🧪 开始安全调试: {username}")

        if not username:
            return jsonify({'error': '用户名不能为空'}), 400

        if not twitter_api or not twitter_api.client:
            return jsonify({'error': 'Twitter API未初始化'}), 500

        result = {
            'username': username,
            'steps': [],
            'success': False,
            'tweets_found': 0,
            'error': None
        }

        # 步骤1: 获取用户信息（快速测试）
        try:
            logger.info(f"🔍 步骤1: 获取用户信息")
            result['steps'].append('开始获取用户信息')

            user_response = twitter_api.client.get_user(username=username)
            if not user_response or not user_response.data:
                result['error'] = f'用户 {username} 不存在或无法访问'
                result['steps'].append(f'❌ {result["error"]}')
                return jsonify(result)

            user_id = user_response.data.id
            result['user_id'] = str(user_id)
            result['steps'].append(f'✅ 用户ID: {user_id}')
            logger.info(f"✅ 用户ID获取成功: {user_id}")

        except Exception as e:
            result['error'] = f'获取用户信息失败: {str(e)}'
            result['steps'].append(f'❌ {result["error"]}')
            logger.error(f"❌ 获取用户信息失败: {e}")
            return jsonify(result)

        # 步骤2: 尝试最简单的推文获取
        try:
            logger.info(f"🐦 步骤2: 尝试获取推文")
            result['steps'].append('尝试获取推文（最简参数）')

            # 使用最简单的参数，避免复杂的字段请求
            tweets_response = twitter_api.client.get_users_tweets(
                id=user_id,
                max_results=max_results
                # 不添加任何额外字段，使用最基础的调用
            )

            if tweets_response and tweets_response.data:
                tweets_count = len(tweets_response.data)
                result['success'] = True
                result['tweets_found'] = tweets_count
                result['steps'].append(f'✅ 找到 {tweets_count} 条推文')

                # 记录第一条推文信息
                if tweets_response.data:
                    first_tweet = tweets_response.data[0]
                    result['first_tweet'] = {
                        'id': str(first_tweet.id),
                        'text': first_tweet.text[:100] if first_tweet.text else '',
                        'created_at': str(first_tweet.created_at) if hasattr(first_tweet, 'created_at') else None
                    }
                    result['steps'].append(f'📝 样本: {result["first_tweet"]["text"][:50]}...')

                logger.info(f"✅ 推文获取成功: {tweets_count} 条")

            else:
                result['steps'].append('⚠️ API响应为空，未找到推文')
                logger.warning(f"⚠️ API响应为空")

        except Exception as e:
            result['error'] = f'推文获取失败: {str(e)}'
            result['steps'].append(f'❌ {result["error"]}')
            logger.error(f"❌ 推文获取失败: {e}")

        logger.info(f"🎉 调试完成: {result['success']}")
        return jsonify(result)

    except Exception as e:
        logger.error(f"💥 安全调试发生意外错误: {e}")
        return jsonify({
            'error': f'调试过程异常: {str(e)}',
            'steps': ['调试过程发生意外错误'],
            'success': False
        }), 500

@app.route('/api/quick_test_user', methods=['POST'])
def quick_test_user():
    """快速测试特定用户 - 超轻量版本"""
    try:
        data = request.get_json() or {}
        username = data.get('username', 'karpathy')  # 默认测试karpathy

        logger.info(f"⚡ 快速测试用户: {username}")

        if not twitter_api or not twitter_api.client:
            return jsonify({'error': 'Twitter API未初始化'}), 500

        # 只做最基础的测试
        try:
            # 清理用户名
            clean_username = username.replace('@', '').strip()

            # 获取用户信息
            user_response = twitter_api.client.get_user(username=clean_username)
            if not user_response or not user_response.data:
                return jsonify({
                    'success': False,
                    'message': f'用户 {clean_username} 不存在'
                })

            user_id = user_response.data.id

            # 尝试获取少量推文
            tweets_response = twitter_api.client.get_users_tweets(
                id=user_id,
                max_results=5  # 只获取5条
            )

            success = tweets_response and tweets_response.data and len(tweets_response.data) > 0
            tweets_count = len(tweets_response.data) if tweets_response and tweets_response.data else 0

            return jsonify({
                'success': success,
                'username': clean_username,
                'user_id': str(user_id),
                'tweets_count': tweets_count,
                'message': f'{"✅ 成功" if success else "⚠️ 无推文"} - 用户: {clean_username}, 推文: {tweets_count} 条'
            })

        except Exception as e:
            return jsonify({
                'success': False,
                'error': str(e),
                'message': f'测试失败: {str(e)}'
            })

    except Exception as e:
        logger.error(f"快速测试失败: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/check_system_health', methods=['GET'])
def check_system_health():
    """检查系统健康状态 - 不调用外部API"""
    try:
        health_status = {
            'timestamp': datetime.now().isoformat(),
            'components': {},
            'overall_status': 'healthy'
        }

        # 检查数据库
        try:
            conn = sqlite3.connect('research_platform.db')
            cursor = conn.cursor()
            cursor.execute('SELECT COUNT(*) FROM researchers')
            researchers_count = cursor.fetchone()[0]
            conn.close()

            health_status['components']['database'] = {
                'status': 'healthy',
                'researchers_count': researchers_count
            }
        except Exception as e:
            health_status['components']['database'] = {
                'status': 'error',
                'error': str(e)
            }
            health_status['overall_status'] = 'degraded'

        # 检查Twitter API客户端（不调用API）
        if twitter_api and twitter_api.client:
            health_status['components']['twitter_api'] = {
                'status': 'initialized',
                'api_working': getattr(twitter_api, 'api_working', False),
                'connection_tested': getattr(twitter_api, 'connection_tested', False)
            }
        else:
            health_status['components']['twitter_api'] = {
                'status': 'not_initialized'
            }
            health_status['overall_status'] = 'degraded'

        # 检查监控服务
        if monitoring_service:
            health_status['components']['monitoring'] = {
                'status': 'initialized',
                'running': getattr(monitoring_service, 'running', False)
            }
        else:
            health_status['components']['monitoring'] = {
                'status': 'not_initialized'
            }

        return jsonify(health_status)

    except Exception as e:
        logger.error(f"健康检查失败: {e}")
        return jsonify({
            'overall_status': 'error',
            'error': str(e),
            'timestamp': datetime.now().isoformat()
        }), 500

@app.route('/api/test_api_parameters', methods=['POST'])
def test_api_parameters():
    """测试不同的API参数组合"""
    try:
        data = request.get_json()
        username = data.get('username', '').replace('@', '').strip()
        researcher_id = data.get('researcher_id')

        if not twitter_api or not twitter_api.client:
            return jsonify({'error': 'Twitter API未初始化'}), 500

        # 获取用户ID
        user_response = twitter_api.client.get_user(username=username)
        if not user_response or not user_response.data:
            return jsonify({'error': f'用户 {username} 不存在'}), 404

        user_id = user_response.data.id

        # 定义不同的参数测试组合
        test_combinations = [
            {
                'name': '最简参数',
                'params': {'id': user_id, 'max_results': 10}
            },
            {
                'name': '基础字段',
                'params': {
                    'id': user_id,
                    'max_results': 10,
                    'tweet_fields': ['created_at']
                }
            },
            {
                'name': '公共指标',
                'params': {
                    'id': user_id,
                    'max_results': 10,
                    'tweet_fields': ['created_at', 'public_metrics']
                }
            },
            {
                'name': '排除转发回复',
                'params': {
                    'id': user_id,
                    'max_results': 10,
                    'tweet_fields': ['created_at', 'public_metrics'],
                    'exclude': ['retweets', 'replies']
                }
            },
            {
                'name': '包含转发回复',
                'params': {
                    'id': user_id,
                    'max_results': 10,
                    'tweet_fields': ['created_at', 'public_metrics']
                    # 不设置exclude
                }
            },
            {
                'name': '最近3天',
                'params': {
                    'id': user_id,
                    'max_results': 10,
                    'tweet_fields': ['created_at', 'public_metrics'],
                    'start_time': (datetime.now(timezone.utc) - timedelta(days=3)).isoformat()
                }
            },
            {
                'name': '最近30天',
                'params': {
                    'id': user_id,
                    'max_results': 10,
                    'tweet_fields': ['created_at', 'public_metrics'],
                    'start_time': (datetime.now(timezone.utc) - timedelta(days=30)).isoformat()
                }
            },
            {
                'name': '更大数量',
                'params': {
                    'id': user_id,
                    'max_results': 50,
                    'tweet_fields': ['created_at', 'public_metrics']
                }
            }
        ]

        test_results = []
        any_success = False

        for test in test_combinations:
            try:
                logger.info(f"🧪 测试参数组合: {test['name']}")
                logger.info(f"   参数: {test['params']}")

                response = twitter_api.client.get_users_tweets(**test['params'])

                success = response and response.data and len(response.data) > 0
                if success:
                    any_success = True

                result = {
                    'name': test['name'],
                    'params': test['params'],
                    'success': success,
                    'tweets_count': len(response.data) if response and response.data else 0,
                    'response_info': {
                        'has_response': response is not None,
                        'has_data': response and hasattr(response, 'data') and response.data is not None,
                        'data_type': str(type(response.data)) if response and hasattr(response, 'data') else None
                    }
                }

                if success and response.data:
                    # 记录第一条推文的样本
                    first_tweet = response.data[0]
                    result['sample_tweet'] = {
                        'id': str(first_tweet.id),
                        'text': first_tweet.text[:100] if first_tweet.text else '',
                        'created_at': str(first_tweet.created_at) if hasattr(first_tweet, 'created_at') else None
                    }

                test_results.append(result)
                logger.info(f"   结果: {'✅ 成功' if success else '❌ 失败'} - {result['tweets_count']} 条")

                # 延迟避免速率限制
                time.sleep(1)

            except Exception as e:
                result = {
                    'name': test['name'],
                    'params': test['params'],
                    'success': False,
                    'error': str(e),
                    'error_type': type(e).__name__
                }
                test_results.append(result)
                logger.error(f"   异常: {str(e)}")

        return jsonify({
            'username': username,
            'user_id': str(user_id),
            'tests': test_results,
            'any_success': any_success,
            'successful_tests': [t for t in test_results if t['success']],
            'failed_tests': [t for t in test_results if not t['success']]
        })

    except Exception as e:
        logger.error(f"参数测试失败: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/debug_data_flow', methods=['POST'])
def debug_data_flow():
    """调试数据从API到数据库的完整流程"""
    try:
        data = request.get_json()
        researcher_id = data.get('researcher_id')
        username = data.get('username', '').replace('@', '').strip()

        flow_log = []

        # 步骤1: 检查数据库中现有内容
        flow_log.append("📊 检查数据库现有内容...")

        conn = sqlite3.connect('research_platform.db')
        cursor = conn.cursor()

        cursor.execute('SELECT COUNT(*) FROM x_content WHERE researcher_id = ?', (researcher_id,))
        existing_count = cursor.fetchone()[0]
        flow_log.append(f"   现有内容: {existing_count} 条")

        cursor.execute('SELECT COUNT(*) FROM x_content', ())
        total_count = cursor.fetchone()[0]
        flow_log.append(f"   系统总内容: {total_count} 条")

        # 步骤2: 测试API获取
        flow_log.append("🐦 测试API数据获取...")

        if twitter_api:
            try:
                # 使用get_user_tweets方法
                tweets = twitter_api.get_user_tweets(username, max_results=10)
                flow_log.append(f"   get_user_tweets返回: {len(tweets)} 条")

                if tweets:
                    flow_log.append("   样本推文:")
                    for i, tweet in enumerate(tweets[:2]):
                        flow_log.append(f"     {i+1}. {tweet.get('created_at', 'N/A')} | {tweet.get('content', '')[:50]}...")
                else:
                    flow_log.append("   ❌ 未获取到推文数据")

            except Exception as e:
                flow_log.append(f"   ❌ API调用失败: {str(e)}")
                tweets = []
        else:
            flow_log.append("   ❌ Twitter API未初始化")
            tweets = []

        # 步骤3: 测试数据库插入
        flow_log.append("💾 测试数据库插入...")

        if tweets:
            try:
                # 模拟插入过程（不实际插入）
                insert_test_results = []

                for i, tweet in enumerate(tweets[:3]):  # 只测试前3条
                    try:
                        # 检查是否已存在
                        cursor.execute('SELECT id FROM x_content WHERE tweet_id = ?', (tweet['id'],))
                        existing = cursor.fetchone()

                        if existing:
                            insert_test_results.append(f"推文 {i+1}: 已存在 (ID: {tweet['id']})")
                        else:
                            insert_test_results.append(f"推文 {i+1}: 可插入 (ID: {tweet['id']})")

                    except Exception as e:
                        insert_test_results.append(f"推文 {i+1}: 检查失败 - {str(e)}")

                for result in insert_test_results:
                    flow_log.append(f"   {result}")

            except Exception as e:
                flow_log.append(f"   ❌ 数据库测试失败: {str(e)}")
        else:
            flow_log.append("   ⚠️ 无数据可测试插入")

        # 步骤4: 检查最近的内容
        flow_log.append("🔍 检查最近添加的内容...")

        cursor.execute('''
            SELECT tweet_id, content, created_at, collected_at
            FROM x_content
            WHERE researcher_id = ?
            ORDER BY collected_at DESC
            LIMIT 5
        ''', (researcher_id,))

        recent_content = cursor.fetchall()
        if recent_content:
            flow_log.append(f"   最近 {len(recent_content)} 条内容:")
            for content in recent_content:
                tweet_id, text, created_at, collected_at = content
                flow_log.append(f"     {collected_at} | {created_at} | {text[:50] if text else 'N/A'}...")
        else:
            flow_log.append("   ❌ 该研究者无任何内容")

        conn.close()

        # 步骤5: 系统状态检查
        flow_log.append("⚙️ 系统状态检查...")
        flow_log.append(f"   Twitter API: {'✅ 已初始化' if twitter_api else '❌ 未初始化'}")
        if twitter_api:
            flow_log.append(f"   API状态: {'✅ 正常' if twitter_api.api_working else '❌ 异常'}")
            flow_log.append(f"   连接测试: {'✅ 已测试' if twitter_api.connection_tested else '⚠️ 未测试'}")

        return jsonify({
            'researcher_id': researcher_id,
            'username': username,
            'flow_log': flow_log,
            'summary': {
                'existing_content_count': existing_count,
                'api_tweets_count': len(tweets) if 'tweets' in locals() else 0,
                'api_working': twitter_api and twitter_api.api_working if twitter_api else False,
                'database_accessible': True  # 如果能到这里说明数据库正常
            }
        })

    except Exception as e:
        logger.error(f"数据流程调试失败: {e}")
        return jsonify({
            'error': str(e),
            'flow_log': flow_log if 'flow_log' in locals() else []
        }), 500

@app.route('/api/test_twitter/<int:researcher_id>', methods=['POST'])
def test_twitter_api(researcher_id):
    """测试Twitter API连接和数据获取"""
    try:
        conn = sqlite3.connect('research_platform.db')
        cursor = conn.cursor()

        cursor.execute('SELECT name, x_account FROM researchers WHERE id = ?', (researcher_id,))
        researcher = cursor.fetchone()

        if not researcher:
            conn.close()
            return jsonify({'error': 'Researcher not found'}), 404

        name, x_account = researcher
        conn.close()

        # 测试API连接
        if not twitter_api:
            return jsonify({
                'error': 'Twitter API未初始化',
                'name': name,
                'x_account': x_account
            }), 500

        # 测试获取用户信息
        logger.info(f"🧪 测试获取 {name} ({x_account}) 的用户信息")
        user_info = twitter_api.get_user_info(x_account)

        # 测试获取推文
        logger.info(f"🧪 测试获取 {name} ({x_account}) 的推文")
        tweets = twitter_api.get_user_tweets(x_account, max_results=5)

        return jsonify({
            'message': f'测试完成',
            'name': name,
            'x_account': x_account,
            'api_working': twitter_api.api_working,
            'user_info': user_info,
            'tweets_count': len(tweets) if tweets else 0,
            'tweets_sample': tweets[:2] if tweets else [],  # 返回前2条作为样本
            'test_results': {
                'user_info_success': user_info is not None,
                'tweets_success': tweets is not None and len(tweets) > 0
            }
        })

    except Exception as e:
        logger.error(f"测试Twitter API失败: {e}")
        return jsonify({'error': str(e)}), 500

# 立即在你的app.py文件末尾（在if __name__ == '__main__':之前）添加这个路由：

@app.route('/api/emergency_debug/<int:researcher_id>')
def emergency_debug(researcher_id):
    """紧急调试特定研究者的404问题"""
    try:
        conn = sqlite3.connect('research_platform.db')
        cursor = conn.cursor()
        cursor.execute('SELECT name, x_account FROM researchers WHERE id = ?', (researcher_id,))
        researcher = cursor.fetchone()
        conn.close()

        if not researcher:
            return jsonify({'error': f'研究者ID {researcher_id} 不存在'}), 404

        name, x_account = researcher
        debug_info = {
            'researcher_id': researcher_id,
            'name': name,
            'x_account': x_account,
            'debug_steps': []
        }

        # 步骤1：检查用户名
        debug_info['debug_steps'].append(f"1. 原始用户名: '{x_account}'")

        if not x_account:
            debug_info['debug_steps'].append("❌ 用户名为空 - 这是问题所在！")
            debug_info['issue'] = 'empty_username'
            return jsonify(debug_info)

        clean_username = x_account.replace('@', '').strip()
        debug_info['clean_username'] = clean_username
        debug_info['debug_steps'].append(f"2. 清理后用户名: '{clean_username}'")

        if not clean_username:
            debug_info['debug_steps'].append("❌ 清理后用户名为空 - 这是问题所在！")
            debug_info['issue'] = 'invalid_username'
            return jsonify(debug_info)

        # 步骤2：检查API状态
        debug_info['debug_steps'].append(f"3. Twitter API状态:")
        debug_info['debug_steps'].append(f"   - 客户端存在: {twitter_api and twitter_api.client is not None}")
        debug_info['debug_steps'].append(f"   - API工作状态: {twitter_api and twitter_api.api_working}")

        if not twitter_api or not twitter_api.client:
            debug_info['debug_steps'].append("❌ Twitter API未初始化")
            debug_info['issue'] = 'api_not_initialized'
            return jsonify(debug_info)

        # 步骤3：尝试API调用
        debug_info['debug_steps'].append(f"4. 尝试API调用用户: @{clean_username}")

        try:
            user_response = twitter_api.client.get_user(username=clean_username)

            if not user_response:
                debug_info['debug_steps'].append("❌ API响应为空")
                debug_info['issue'] = 'empty_response'
            elif not user_response.data:
                debug_info['debug_steps'].append("❌ 用户数据为空 - 用户不存在")
                debug_info['issue'] = 'user_not_found'
            else:
                debug_info['debug_steps'].append("✅ 用户存在!")
                debug_info['user_info'] = {
                    'id': str(user_response.data.id),
                    'username': user_response.data.username,
                    'name': user_response.data.name
                }
                debug_info['issue'] = 'none'

        except tweepy.NotFound as e:
            debug_info['debug_steps'].append(f"❌ 404错误 - 用户不存在: {str(e)}")
            debug_info['issue'] = 'user_not_found'
            debug_info['error_details'] = str(e)

        except tweepy.Unauthorized as e:
            debug_info['debug_steps'].append(f"❌ 401错误 - 无权访问: {str(e)}")
            debug_info['issue'] = 'unauthorized'
            debug_info['error_details'] = str(e)

        except Exception as e:
            debug_info['debug_steps'].append(f"❌ 其他错误: {type(e).__name__}: {str(e)}")
            debug_info['issue'] = 'other_error'
            debug_info['error_details'] = str(e)

        # 步骤4：提供解决建议
        if debug_info.get('issue') == 'user_not_found':
            debug_info['suggestions'] = [
                f"用户 @{clean_username} 不存在",
                "可能的原因：1) 用户名错误 2) 账户已删除 3) 账户已暂停",
                "建议：在Twitter上手动搜索该用户名验证"
            ]
        elif debug_info.get('issue') == 'unauthorized':
            debug_info['suggestions'] = [
                f"用户 @{clean_username} 存在但无法访问",
                "可能的原因：1) 私人账户 2) API权限不足",
                "建议：检查该账户是否设为私人"
            ]

        return jsonify(debug_info)

    except Exception as e:
        return jsonify({
            'error': f'调试过程出错: {str(e)}',
            'researcher_id': researcher_id
        }), 500

# 同时添加一个快速修复端点：
@app.route('/api/quick_fix_username/<int:researcher_id>', methods=['POST'])
def quick_fix_username(researcher_id):
    """快速修复用户名格式问题"""
    try:
        data = request.get_json() or {}
        new_username = data.get('new_username', '').strip()

        conn = sqlite3.connect('research_platform.db')
        cursor = conn.cursor()

        if new_username:
            # 确保用户名格式正确
            if not new_username.startswith('@'):
                new_username = '@' + new_username

            cursor.execute('UPDATE researchers SET x_account = ? WHERE id = ?', (new_username, researcher_id))
            conn.commit()

            result = {'message': f'用户名已更新为: {new_username}'}
        else:
            # 获取当前信息
            cursor.execute('SELECT name, x_account FROM researchers WHERE id = ?', (researcher_id,))
            researcher = cursor.fetchone()
            result = {
                'current_name': researcher[0] if researcher else 'Unknown',
                'current_username': researcher[1] if researcher else 'None',
                'message': '请提供新的用户名'
            }

        conn.close()
        return jsonify(result)

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/raw_test')
def raw_test():
    import tweepy
    import os

    token = os.environ.get('TWITTER_BEARER_TOKEN')
    if not token:
        return {'error': 'Token未配置'}

    try:
        client = tweepy.Client(bearer_token=token)
        user = client.get_user(username='karpathy')

        if user and user.data:
            return {
                'success': True,
                'user_id': str(user.data.id),
                'username': user.data.username,
                'name': user.data.name
            }
        else:
            return {'error': 'API响应为空'}

    except Exception as e:
        return {'error': f'{type(e).__name__}: {str(e)}'}

if __name__ == '__main__':
    logger.info("🚀 AI研究者X内容学习平台启动中...")
    logger.info(f"📊 系统容量: 最大支持 5000 位研究者监控")
    logger.info(f"💾 数据库文件: research_platform.db")

    # 检查数据库状态
    if researcher_manager:
        import os
        if os.path.exists('research_platform.db'):
            conn = sqlite3.connect('research_platform.db')
            cursor = conn.cursor()
            cursor.execute('SELECT COUNT(*) FROM researchers')
            count = cursor.fetchone()[0]
            conn.close()
            logger.info(f"📋 数据库状态: 已有 {count} 位研究者")
        else:
            logger.info("📋 数据库状态: 新建数据库")

    logger.info(f"🔑 Twitter API: {'✅ 已配置' if TWITTER_BEARER_TOKEN else '⚠️ 未配置，将无法获取真实数据'}")

    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port, debug=False)
